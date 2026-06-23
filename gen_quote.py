#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《数学之美 软件授权报价单》Word 文档(专业排版)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CN_FONT = '微软雅黑'
EN_FONT = 'Calibri'


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_font(run, size=None, bold=None, color=None, cn=CN_FONT, en=EN_FONT):
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cn)


def add_heading(doc, text, size=15, color=NAVY, space_before=16, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '4F46E5')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def add_body(doc, text, size=10.5, color=GRAY, bold=False, align=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.45
    if align: p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=10.5, color=GRAY):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_table(doc, headers, rows, widths, header_bg='1F3A5F', zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Cm(widths[i]); set_cell_bg(c, header_bg)
        pr = c.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pr.add_run(h), size=10, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].width = Cm(widths[ci])
            if zebra and ri % 2 == 1: set_cell_bg(cells[ci], 'F2F4F8')
            pr = cells[ci].paragraphs[0]
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(pr.add_run(val), size=9.5, color=GRAY)
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)
normal = doc.styles['Normal']
normal.font.name = EN_FONT
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)

# ===== 封面 =====
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('「 数 学 之 美 」'), size=30, bold=True, color=NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
set_run_font(p.add_run('交互式数学可视化实验平台'), size=16, bold=True, color=INDIGO)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('Interactive Mathematics Visualization Platform'), size=11, color=LIGHT)
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run('软  件  授  权  报  价  单'), size=20, bold=True, color=GRAY)
for _ in range(4):
    doc.add_paragraph()

info = [
    ('文档名称', '软件授权报价单'),
    ('产品版本', '数学之美 v2.0'),
    ('授权方式', '双授权(非商业免费 / 商业授权)'),
    ('适用对象', '教育机构 / 在线学习平台'),
    ('报价有效期', '自报价之日起 30 个自然日'),
    ('编制日期', '2026 年 6 月'),
]
it = doc.add_table(rows=len(info), cols=2)
it.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    kc = it.rows[i].cells[0]; vc = it.rows[i].cells[1]
    kc.width = Cm(4.5); vc.width = Cm(9)
    set_cell_bg(kc, 'EEF1F7')
    pk = kc.paragraphs[0]; pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(pk.add_run(k), size=10.5, bold=True, color=NAVY)
    pv = vc.paragraphs[0]
    set_run_font(pv.add_run(v), size=10.5, color=GRAY)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ===== 一、产品概述 =====
add_heading(doc, '一、产品概述')
add_body(doc, '「数学之美」是一套面向教育场景的交互式数学可视化实验平台，覆盖从基础算术到高等数学、'
              '概率统计、线性代数、微分方程、混沌与分形等 14 个学科门类。每个实验均为可交互的'
              '动态可视化场景，并配有专业男女双声中文语音讲解，帮助学生在“看得见、调得动、听得懂”'
              '的环境中理解抽象数学概念。平台采用纯前端架构，可独立部署于机构内网或公有云，无需后端依赖。')

add_heading(doc, '二、核心资产清单')
add_table(doc,
    ['资产项', '规模 / 说明'],
    [
        ['交互式数学实验场景', '55 个，覆盖 14 个学科门类'],
        ['语音讲解音频', '5187 个 mp3（男声云希 / 女声晓晓 双声完整覆盖）'],
        ['源代码规模', '约 92,000 行 TypeScript / TSX'],
        ['技术栈', 'React 19 + Vite 7 + TypeScript + Plotly.js 3 + KaTeX + mathjs'],
        ['真仿真内核', '元胞自动机、反应扩散、N 体引力、傅里叶合成等真实数值计算'],
        ['部署形态', '纯静态前端，支持内网 / 公有云独立部署'],
        ['第三方依赖协议', '均为 MIT / BSD / Apache 等宽松开源协议，允许商用'],
    ],
    widths=[5.5, 11.5])
add_body(doc, '说明：以上资产若按市场人天成本从零重建，预估投入约 30–60 万元人民币，且需 8–12 个月研发周期。',
         size=9.5, color=LIGHT)

# ===== 三、授权模式说明 =====
add_heading(doc, '三、授权模式：双授权')
add_body(doc, '本平台以双授权（Dual Licensing）模式发布，公开版本依据 PolyForm Noncommercial 1.0.0 协议授权。'
              '在该协议下，个人学习研究，以及学校、公益研究组织、政府机构等非商业组织的非商业使用，'
              '均完全免费。但任何商业用途——包括企业在生产环境部署、以付费产品 / SaaS / 培训课程形式'
              '对外提供服务、集成进商业系统等——都必须购买【商业授权】。',
         color=GRAY)
add_body(doc, '出品方完整保留软件著作权，公开发布不影响以商业条款另行授权的权利。',
         size=9.5, color=LIGHT)

add_heading(doc, '四、商业授权方案与报价', size=14)
add_body(doc, '面向教育机构提供以下三档商业授权方案。推荐【标准版】，含完整源码访问与一年技术支持。',
         size=10.5, color=GRAY)
add_table(doc,
    ['授权版本', '价格区间', '商业授权内容', '适用场景'],
    [
        ['基础版', '28–35 万', '单机构商业授权，含部署不含定制', '直接上线使用'],
        ['标准版（推荐）', '38–48 万', '商业授权 + 完整源码访问 + 非独家 + 1 年支持', '需二次开发/自主维护'],
        ['旗舰版', '55–68 万', '商业授权 + 定制开发 + 2 年支持 + 品牌 OEM', '深度集成/贴牌运营'],
    ],
    widths=[3.2, 3.0, 6.8, 4.0])
add_body(doc, '注：标准版报价锚定 45 万元，最低成交价不低于 30 万元。以上均为人民币含税价。',
         size=9.5, color=LIGHT)

add_heading(doc, '五、标准版报价明细', size=14)
add_table(doc,
    ['项目', '内容', '金额（元）'],
    [
        ['商业授权费', '商业场景使用权，含 55 场景源码访问 + 5187 音频资产，永久、非独家', '380,000'],
        ['部署与培训', '环境部署、2 个工作日现场/远程培训', '20,000'],
        ['首年技术支持', '远程答疑、缺陷修复、小版本升级', '30,000'],
        ['合计', '—', '430,000'],
    ],
    widths=[4.5, 9.0, 3.5])
add_body(doc, '付款方式：签约预付 50%，源码与资产交付付 30%，验收通过后付 20%。',
         size=10, bold=True, color=NAVY)

# ===== 六、增值服务 =====
add_heading(doc, '六、增值服务（可选）')
add_table(doc,
    ['服务项', '单价', '说明'],
    [
        ['新增定制实验场景', '1.0–1.5 万 / 个', '含交互可视化 + 双声语音讲解'],
        ['年度维护服务', '授权费的 15% / 年', '次年起，含升级与技术支持'],
        ['品牌 OEM 定制', '1.5 万 / 人天', 'Logo、配色、域名等贴牌改造'],
        ['额外技术支持', '1.5–2.0 万 / 人天', '超出约定范围的现场支持'],
    ],
    widths=[4.5, 4.0, 8.5])

# ===== 七、合同关键条款 =====
add_heading(doc, '七、合同关键条款')
add_bullet(doc, '授权性质：非独家商业授权。买方获得在商业场景下使用本软件的权利；卖方保留著作权，可继续向其他客户授权。')
add_bullet(doc, '与公开版本的关系：本软件同时以 PolyForm Noncommercial 协议公开发布（仅限非商业使用），商业授权不影响该公开版本的持续发布。')
add_bullet(doc, '源码使用：买方可自行二次开发、部署、运营，但不得将源代码转售、转授权或单独公开发布。')
add_bullet(doc, '知识产权：软件原始著作权归卖方所有；第三方开源组件遵循其各自的开源协议（MIT/BSD/Apache 等）。')
add_bullet(doc, '验收标准：以 55 个实验场景可正常运行、双声语音正常播放为验收依据，验收期 10 个工作日。')
add_bullet(doc, '技术支持边界：首年支持含缺陷修复与小版本升级；定制开发、新增场景按增值服务单独计费。')
add_bullet(doc, '保密义务：双方对商业授权协议、商务条款、技术资料互负保密义务。')

# ===== 八、联系方式 =====
add_heading(doc, '八、联系方式')
add_body(doc, '如对本报价有任何疑问或需调整授权范围，欢迎随时联系商务洽谈。')
add_body(doc, '商务联系：Zhang Zhen        手机：13501033928')
add_body(doc, '电子邮箱：zhangzhen@gmail.com        签署日期：________________________')
add_body(doc, '本报价单最终解释权归出品方所有。', size=9, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('数学之美_软件授权报价单.docx')
print('base ok')
